import os
from docx import Document
from ebooklib import epub
from lxml import etree

def create_samples():
    base_dir = r"c:\Users\jorge\Documents\babel\input"
    if not os.path.exists(base_dir):
        os.makedirs(base_dir)

    # 1. Create Sample CXML
    root = etree.Element("root")
    child1 = etree.SubElement(root, "p")
    child1.text = "Hello world from CXML."
    child2 = etree.SubElement(root, "note")
    child2.text = "This is a test note."
    tree = etree.ElementTree(root)
    tree.write(os.path.join(base_dir, "sample.cxml"), encoding='utf-8', xml_declaration=True)

    # 2. Create Sample DOCX
    doc = Document()
    doc.add_heading('Sample Document', 0)
    doc.add_paragraph('This is the first paragraph of the DOCX file.')
    doc.add_paragraph('This is the second paragraph with some bold text.', style='List Bullet')
    doc.save(os.path.join(base_dir, "sample.docx"))

    # 3. Create Sample EPUB
    book = epub.EpubBook()
    book.set_identifier('id123456')
    book.set_title('Sample Epub')
    book.set_language('en')
    
    c1 = epub.EpubHtml(title='Intro', file_name='intro.xhtml', lang='en')
    c1.content = '<html><head></head><body><h1>Introduction</h1><p>This is the EPUB content.</p></body></html>'
    book.add_item(c1)
    
    # Add NCX and Nav
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    book.spine = ['nav', c1] 
    epub.write_epub(os.path.join(base_dir, "sample.epub"), book, {})

    print("Samples created.")

if __name__ == "__main__":
    create_samples()
