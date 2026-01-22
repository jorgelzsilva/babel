from docx import Document
from .base import BaseHandler

class DocxHandler(BaseHandler):
    def read_file(self):
        self.doc = Document(self.file_path)

    def extract_text(self):
        """
        Extracts paragraphs as text nodes.
        Limitation: Replacing paragraph.text removes run-level formatting (bold words),
        but keeps paragraph-level styles.
        """
        text_nodes = []
        for para in self.doc.paragraphs:
            if para.text and para.text.strip():
                # We wrap the paragraph in a simple object/adapter if needed, 
                # or just return the para itself since it has .text property
                text_nodes.append(para)
                
        # Also handle tables?
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text and para.text.strip():
                            text_nodes.append(para)

        return text_nodes

    def apply_translations(self, original_texts, translated_texts):
        pass

    def save_file(self, output_path):
        self.doc.save(output_path)
